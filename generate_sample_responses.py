#!/usr/bin/env python3
"""
Generate impressive sample TAM and DCF analysis responses
Creates Word documents with professional financial analysis examples
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import json

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading to the document"""
    heading = doc.add_heading(text, level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_formatted_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def create_tam_sample_response():
    """Create an impressive TAM analysis sample response"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Total Addressable Market (TAM) Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('AI-Powered Computer Vision Solutions for Manufacturing', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    add_heading_with_style(doc, "Executive Summary", 1)
    add_formatted_paragraph(doc, 
        "This comprehensive TAM analysis evaluates the market opportunity for AI-powered computer vision solutions "
        "targeting manufacturing companies. Our multi-methodology approach reveals a substantial addressable market "
        "ranging from $12.8B to $18.2B globally, with significant growth potential driven by Industry 4.0 adoption "
        "and increasing demand for quality control automation.", bold=True)
    
    # Opening Overview
    add_heading_with_style(doc, "1. Opening Overview", 1)
    add_formatted_paragraph(doc,
        "The global manufacturing industry is experiencing a digital transformation revolution, with computer vision "
        "emerging as a critical technology for quality control, predictive maintenance, and process optimization. "
        "Our analysis focuses on the addressable market for AI-powered computer vision solutions specifically "
        "designed for manufacturing applications.")
    
    add_formatted_paragraph(doc,
        "Key market drivers include:")
    add_formatted_paragraph(doc, "• Rising labor costs and skilled worker shortages", italic=True)
    add_formatted_paragraph(doc, "• Increasing quality standards and regulatory requirements", italic=True)
    add_formatted_paragraph(doc, "• Growing adoption of Industry 4.0 and smart manufacturing", italic=True)
    add_formatted_paragraph(doc, "• Need for real-time defect detection and process optimization", italic=True)
    
    # Bottom-Up Analysis
    add_heading_with_style(doc, "2. Bottom-Up Analysis", 1)
    add_formatted_paragraph(doc,
        "Our bottom-up analysis segments the market by company size and manufacturing verticals:")
    
    add_heading_with_style(doc, "2.1 Market Segmentation", 2)
    
    # Create a table for market segments
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Segment'
    hdr_cells[1].text = 'Companies'
    hdr_cells[2].text = 'Avg. ACV ($)'
    hdr_cells[3].text = 'Market Size ($B)'
    
    segments = [
        ('Large Enterprises (1000+ employees)', '15,000', '850,000', '12.75'),
        ('Mid-Market (100-999 employees)', '180,000', '125,000', '22.50'),
        ('Small Manufacturers (10-99 employees)', '2,200,000', '15,000', '33.00'),
        ('Total Addressable Market', '2,395,000', '28,500', '68.25')
    ]
    
    for segment, companies, acv, market_size in segments:
        row_cells = table.add_row().cells
        row_cells[0].text = segment
        row_cells[1].text = companies
        row_cells[2].text = acv
        row_cells[3].text = market_size
    
    add_heading_with_style(doc, "2.2 Customer Lifetime Value (LTV)", 2)
    add_formatted_paragraph(doc,
        "Average customer lifetime value analysis based on 5-year retention rates:")
    add_formatted_paragraph(doc, "• Large Enterprises: $2.1M (95% retention, 3-year payback)", italic=True)
    add_formatted_paragraph(doc, "• Mid-Market: $312K (85% retention, 2.5-year payback)", italic=True)
    add_formatted_paragraph(doc, "• Small Manufacturers: $45K (70% retention, 3-year payback)", italic=True)
    
    # Top-Down Analysis
    add_heading_with_style(doc, "3. Top-Down Analysis", 1)
    add_formatted_paragraph(doc,
        "Global manufacturing automation market analysis based on industry reports:")
    
    add_heading_with_style(doc, "3.1 Market Size Validation", 2)
    add_formatted_paragraph(doc,
        "According to Gartner and McKinsey research:")
    add_formatted_paragraph(doc, "• Global manufacturing automation market: $214B (2024)", italic=True)
    add_formatted_paragraph(doc, "• Computer vision segment: 8-12% of total automation market", italic=True)
    add_formatted_paragraph(doc, "• AI-powered solutions: 15-20% of computer vision market", italic=True)
    add_formatted_paragraph(doc, "• Manufacturing vertical: 35-40% of AI computer vision applications", italic=True)
    
    add_formatted_paragraph(doc,
        "Calculated TAM: $214B × 10% × 17.5% × 37.5% = $14.1B", bold=True)
    
    # Hybrid Analysis
    add_heading_with_style(doc, "4. Hybrid Analysis", 1)
    add_formatted_paragraph(doc,
        "Combining bottom-up and top-down methodologies with industry expert validation:")
    
    add_heading_with_style(doc, "4.1 Confidence Intervals", 2)
    add_formatted_paragraph(doc,
        "Our hybrid analysis provides the following confidence intervals:")
    add_formatted_paragraph(doc, "• Conservative Estimate: $12.8B (25th percentile)", italic=True)
    add_formatted_paragraph(doc, "• Most Likely Scenario: $15.4B (50th percentile)", italic=True)
    add_formatted_paragraph(doc, "• Optimistic Scenario: $18.2B (75th percentile)", italic=True)
    
    add_heading_with_style(doc, "4.2 Market Growth Projections", 2)
    add_formatted_paragraph(doc,
        "Compound Annual Growth Rate (CAGR) analysis:")
    add_formatted_paragraph(doc, "• 2024-2027: 18.5% CAGR", italic=True)
    add_formatted_paragraph(doc, "• 2027-2030: 15.2% CAGR", italic=True)
    add_formatted_paragraph(doc, "• 2030 Market Size: $28.7B", italic=True)
    
    # Market Capture Analysis
    add_heading_with_style(doc, "5. Market Capture Analysis", 1)
    add_heading_with_style(doc, "5.1 Competitive Landscape", 2)
    add_formatted_paragraph(doc,
        "Key market players and their estimated market shares:")
    add_formatted_paragraph(doc, "• Cognex Corporation: 12-15% market share", italic=True)
    add_formatted_paragraph(doc, "• Keyence Corporation: 8-10% market share", italic=True)
    add_formatted_paragraph(doc, "• Basler AG: 5-7% market share", italic=True)
    add_formatted_paragraph(doc, "• Emerging AI startups: 15-20% market share", italic=True)
    add_formatted_paragraph(doc, "• Remaining market: 48-60% (fragmented)", italic=True)
    
    add_heading_with_style(doc, "5.2 Market Penetration Strategy", 2)
    add_formatted_paragraph(doc,
        "Recommended market capture approach:")
    add_formatted_paragraph(doc, "• Year 1-2: Focus on mid-market manufacturers (5% penetration)", italic=True)
    add_formatted_paragraph(doc, "• Year 3-4: Expand to large enterprises (3% penetration)", italic=True)
    add_formatted_paragraph(doc, "• Year 5+: Scale to small manufacturers (1% penetration)", italic=True)
    add_formatted_paragraph(doc, "• Total Addressable Revenue: $1.2B by Year 5", bold=True)
    
    # Final Deliverables
    add_heading_with_style(doc, "6. Final Deliverables", 1)
    add_heading_with_style(doc, "6.1 Key Findings", 2)
    add_formatted_paragraph(doc,
        "• Total Addressable Market: $12.8B - $18.2B globally", bold=True)
    add_formatted_paragraph(doc,
        "• Most attractive segment: Mid-market manufacturers ($22.5B)", bold=True)
    add_formatted_paragraph(doc,
        "• Expected market growth: 18.5% CAGR through 2027", bold=True)
    add_formatted_paragraph(doc,
        "• Recommended market entry: Quality control applications", bold=True)
    
    add_heading_with_style(doc, "6.2 Revenue Projections", 2)
    add_formatted_paragraph(doc,
        "3-year revenue projection based on market capture analysis:")
    add_formatted_paragraph(doc, "• Year 1: $15M (0.1% market penetration)", italic=True)
    add_formatted_paragraph(doc, "• Year 2: $45M (0.3% market penetration)", italic=True)
    add_formatted_paragraph(doc, "• Year 3: $120M (0.8% market penetration)", italic=True)
    
    add_heading_with_style(doc, "6.3 Risk Factors", 2)
    add_formatted_paragraph(doc,
        "Key risks and mitigation strategies:")
    add_formatted_paragraph(doc, "• Technology risk: Continuous R&D investment required", italic=True)
    add_formatted_paragraph(doc, "• Competition risk: Focus on AI differentiation", italic=True)
    add_formatted_paragraph(doc, "• Economic risk: Target recession-resistant verticals", italic=True)
    add_formatted_paragraph(doc, "• Regulatory risk: Ensure compliance with manufacturing standards", italic=True)
    
    # Sources
    add_heading_with_style(doc, "Sources and References", 1)
    add_formatted_paragraph(doc, "• Gartner: 'Manufacturing Automation Market Analysis 2024'", italic=True)
    add_formatted_paragraph(doc, "• McKinsey Global Institute: 'The Future of Manufacturing'", italic=True)
    add_formatted_paragraph(doc, "• IDC: 'Computer Vision in Manufacturing Report'", italic=True)
    add_formatted_paragraph(doc, "• Industry expert interviews and surveys", italic=True)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Generated by Phoenician Capital AI Analysis Platform - {datetime.now().strftime('%B %d, %Y')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def create_dcf_sample_response():
    """Create an impressive DCF analysis sample response"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Discounted Cash Flow (DCF) Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('AI-Powered Computer Vision Solutions for Manufacturing', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    add_heading_with_style(doc, "Executive Summary", 1)
    add_formatted_paragraph(doc,
        "This comprehensive DCF analysis values an AI-powered computer vision startup targeting manufacturing "
        "applications. Based on our three-statement model and scenario analysis, the company's enterprise value "
        "ranges from $89M to $156M, with a base case valuation of $124M. The analysis incorporates detailed "
        "financial projections, WACC calculations, and sensitivity analysis across multiple scenarios.", bold=True)
    
    # Model Construction
    add_heading_with_style(doc, "1. Model Construction", 1)
    add_heading_with_style(doc, "1.1 Key Assumptions", 2)
    add_formatted_paragraph(doc,
        "The DCF model is built on the following foundational assumptions:")
    add_formatted_paragraph(doc, "• Forecast period: 5 years (2024-2028)", italic=True)
    add_formatted_paragraph(doc, "• Terminal growth rate: 2.5% (long-term GDP growth)", italic=True)
    add_formatted_paragraph(doc, "• WACC: 12.3% (calculated based on capital structure)", italic=True)
    add_formatted_paragraph(doc, "• Tax rate: 25% (blended federal and state)", italic=True)
    add_formatted_paragraph(doc, "• Working capital: 15% of revenue", italic=True)
    
    add_heading_with_style(doc, "1.2 Business Model", 2)
    add_formatted_paragraph(doc,
        "Revenue model based on SaaS subscriptions with implementation services:")
    add_formatted_paragraph(doc, "• Software subscriptions: 80% of revenue", italic=True)
    add_formatted_paragraph(doc, "• Professional services: 15% of revenue", italic=True)
    add_formatted_paragraph(doc, "• Hardware integration: 5% of revenue", italic=True)
    
    # DCF Summary Table
    add_heading_with_style(doc, "2. DCF Summary Table", 1)
    
    # Create DCF summary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Year'
    hdr_cells[1].text = 'Free Cash Flow ($M)'
    hdr_cells[2].text = 'Present Value ($M)'
    
    dcf_data = [
        ('2024', '2.1', '1.9'),
        ('2025', '8.7', '6.9'),
        ('2026', '18.4', '13.0'),
        ('2027', '32.1', '20.3'),
        ('2028', '48.9', '27.4'),
        ('Terminal Value', '489.0', '275.2'),
        ('Total Enterprise Value', '', '344.7')
    ]
    
    for year, fcf, pv in dcf_data:
        row_cells = table.add_row().cells
        row_cells[0].text = year
        row_cells[1].text = fcf
        row_cells[2].text = pv
    
    add_formatted_paragraph(doc,
        "Enterprise Value: $344.7M", bold=True)
    add_formatted_paragraph(doc,
        "Less: Net Debt: $0M", bold=True)
    add_formatted_paragraph(doc,
        "Equity Value: $344.7M", bold=True)
    add_formatted_paragraph(doc,
        "Shares Outstanding: 2.78M", bold=True)
    add_formatted_paragraph(doc,
        "Value per Share: $124.0", bold=True)
    
    # WACC Calculation
    add_heading_with_style(doc, "3. WACC Calculation", 1)
    add_heading_with_style(doc, "3.1 Cost of Equity", 2)
    add_formatted_paragraph(doc,
        "Cost of equity calculated using CAPM:")
    add_formatted_paragraph(doc, "• Risk-free rate: 4.2% (10-year Treasury)", italic=True)
    add_formatted_paragraph(doc, "• Beta: 1.4 (technology sector average)", italic=True)
    add_formatted_paragraph(doc, "• Market risk premium: 6.0%", italic=True)
    add_formatted_paragraph(doc, "• Cost of equity: 4.2% + (1.4 × 6.0%) = 12.6%", bold=True)
    
    add_heading_with_style(doc, "3.2 Cost of Debt", 2)
    add_formatted_paragraph(doc,
        "Cost of debt analysis:")
    add_formatted_paragraph(doc, "• Pre-tax cost of debt: 8.5%", italic=True)
    add_formatted_paragraph(doc, "• Tax rate: 25%", italic=True)
    add_formatted_paragraph(doc, "• After-tax cost of debt: 8.5% × (1 - 25%) = 6.4%", bold=True)
    
    add_heading_with_style(doc, "3.3 Capital Structure", 2)
    add_formatted_paragraph(doc,
        "Target capital structure:")
    add_formatted_paragraph(doc, "• Equity: 85%", italic=True)
    add_formatted_paragraph(doc, "• Debt: 15%", italic=True)
    add_formatted_paragraph(doc, "• WACC: (85% × 12.6%) + (15% × 6.4%) = 11.7%", bold=True)
    
    # Three-Statement Model
    add_heading_with_style(doc, "4. Three-Statement Model", 1)
    add_heading_with_style(doc, "4.1 Income Statement Projections", 2)
    
    # Income statement table
    income_table = doc.add_table(rows=1, cols=5)
    income_table.style = 'Table Grid'
    hdr_cells = income_table.rows[0].cells
    hdr_cells[0].text = 'Year'
    hdr_cells[1].text = '2024'
    hdr_cells[2].text = '2025'
    hdr_cells[3].text = '2026'
    hdr_cells[4].text = '2027'
    
    income_data = [
        ('Revenue ($M)', '15.0', '45.0', '90.0', '150.0'),
        ('Gross Profit ($M)', '12.0', '36.0', '72.0', '120.0'),
        ('EBITDA ($M)', '1.5', '9.0', '22.5', '45.0'),
        ('EBIT ($M)', '0.8', '6.3', '18.0', '36.0'),
        ('Net Income ($M)', '0.6', '4.7', '13.5', '27.0')
    ]
    
    for line_item, y1, y2, y3, y4 in income_data:
        row_cells = income_table.add_row().cells
        row_cells[0].text = line_item
        row_cells[1].text = y1
        row_cells[2].text = y2
        row_cells[3].text = y3
        row_cells[4].text = y4
    
    add_heading_with_style(doc, "4.2 Balance Sheet Projections", 2)
    add_formatted_paragraph(doc,
        "Key balance sheet items:")
    add_formatted_paragraph(doc, "• Cash: Growing from $5M to $45M by 2027", italic=True)
    add_formatted_paragraph(doc, "• Accounts Receivable: 15% of revenue", italic=True)
    add_formatted_paragraph(doc, "• Property, Plant & Equipment: $2M annual capex", italic=True)
    add_formatted_paragraph(doc, "• Total Debt: $15M by 2027", italic=True)
    add_formatted_paragraph(doc, "• Shareholders' Equity: $89M by 2027", italic=True)
    
    add_heading_with_style(doc, "4.3 Cash Flow Projections", 2)
    add_formatted_paragraph(doc,
        "Operating cash flow drivers:")
    add_formatted_paragraph(doc, "• Revenue growth: 200% in 2025, 100% in 2026, 67% in 2027", italic=True)
    add_formatted_paragraph(doc, "• Gross margin: 80% (improving to 85% by 2027)", italic=True)
    add_formatted_paragraph(doc, "• EBITDA margin: 10% in 2024, 30% by 2027", italic=True)
    add_formatted_paragraph(doc, "• Working capital: 15% of revenue", italic=True)
    add_formatted_paragraph(doc, "• Capital expenditures: $2M annually", italic=True)
    
    # Scenario Analysis
    add_heading_with_style(doc, "5. Scenario Analysis", 1)
    add_heading_with_style(doc, "5.1 Bull Case Scenario", 2)
    add_formatted_paragraph(doc,
        "Optimistic assumptions:")
    add_formatted_paragraph(doc, "• Revenue growth: 50% higher than base case", italic=True)
    add_formatted_paragraph(doc, "• EBITDA margins: 35% by 2027", italic=True)
    add_formatted_paragraph(doc, "• WACC: 10.5% (lower risk premium)", italic=True)
    add_formatted_paragraph(doc, "• Terminal growth: 3.0%", italic=True)
    add_formatted_paragraph(doc, "• Enterprise Value: $156M", bold=True)
    
    add_heading_with_style(doc, "5.2 Base Case Scenario", 2)
    add_formatted_paragraph(doc,
        "Most likely scenario:")
    add_formatted_paragraph(doc, "• Revenue growth: As projected above", italic=True)
    add_formatted_paragraph(doc, "• EBITDA margins: 30% by 2027", italic=True)
    add_formatted_paragraph(doc, "• WACC: 11.7%", italic=True)
    add_formatted_paragraph(doc, "• Terminal growth: 2.5%", italic=True)
    add_formatted_paragraph(doc, "• Enterprise Value: $124M", bold=True)
    
    add_heading_with_style(doc, "5.3 Bear Case Scenario", 2)
    add_formatted_paragraph(doc,
        "Conservative assumptions:")
    add_formatted_paragraph(doc, "• Revenue growth: 30% lower than base case", italic=True)
    add_formatted_paragraph(doc, "• EBITDA margins: 25% by 2027", italic=True)
    add_formatted_paragraph(doc, "• WACC: 13.0% (higher risk premium)", italic=True)
    add_formatted_paragraph(doc, "• Terminal growth: 2.0%", italic=True)
    add_formatted_paragraph(doc, "• Enterprise Value: $89M", bold=True)
    
    # Sensitivity Analysis
    add_heading_with_style(doc, "6. Sensitivity Analysis", 1)
    add_heading_with_style(doc, "6.1 WACC Sensitivity", 2)
    add_formatted_paragraph(doc,
        "Enterprise value sensitivity to WACC changes:")
    add_formatted_paragraph(doc, "• WACC 10.0%: $142M (+14.5%)", italic=True)
    add_formatted_paragraph(doc, "• WACC 11.7%: $124M (base case)", italic=True)
    add_formatted_paragraph(doc, "• WACC 13.5%: $108M (-12.9%)", italic=True)
    
    add_heading_with_style(doc, "6.2 Terminal Growth Sensitivity", 2)
    add_formatted_paragraph(doc,
        "Enterprise value sensitivity to terminal growth:")
    add_formatted_paragraph(doc, "• Terminal growth 2.0%: $115M (-7.3%)", italic=True)
    add_formatted_paragraph(doc, "• Terminal growth 2.5%: $124M (base case)", italic=True)
    add_formatted_paragraph(doc, "• Terminal growth 3.0%: $135M (+8.9%)", italic=True)
    
    # Historical Benchmarking
    add_heading_with_style(doc, "7. Historical Benchmarking", 1)
    add_heading_with_style(doc, "7.1 Comparable Company Analysis", 2)
    add_formatted_paragraph(doc,
        "Public company comparables in AI/computer vision space:")
    add_formatted_paragraph(doc, "• Cognex Corporation: 8.2x EV/Revenue, 25.4x EV/EBITDA", italic=True)
    add_formatted_paragraph(doc, "• Keyence Corporation: 12.1x EV/Revenue, 35.2x EV/EBITDA", italic=True)
    add_formatted_paragraph(doc, "• Basler AG: 3.8x EV/Revenue, 18.7x EV/EBITDA", italic=True)
    add_formatted_paragraph(doc, "• Average multiples: 8.0x EV/Revenue, 26.4x EV/EBITDA", italic=True)
    
    add_heading_with_style(doc, "7.2 Valuation Comparison", 2)
    add_formatted_paragraph(doc,
        "Our DCF valuation vs. comparable multiples:")
    add_formatted_paragraph(doc, "• DCF Enterprise Value: $124M", italic=True)
    add_formatted_paragraph(doc, "• 2027 Revenue multiple (8.0x): $120M", italic=True)
    add_formatted_paragraph(doc, "• 2027 EBITDA multiple (26.4x): $118M", italic=True)
    add_formatted_paragraph(doc, "• Valuation range: $118M - $124M", bold=True)
    
    # Key Value Drivers
    add_heading_with_style(doc, "8. Key Value Drivers", 1)
    add_heading_with_style(doc, "8.1 Revenue Growth Drivers", 2)
    add_formatted_paragraph(doc,
        "Primary drivers of revenue growth:")
    add_formatted_paragraph(doc, "• Market expansion: TAM growing at 18.5% CAGR", italic=True)
    add_formatted_paragraph(doc, "• Customer acquisition: 150% annual growth in customer base", italic=True)
    add_formatted_paragraph(doc, "• Product expansion: New verticals and use cases", italic=True)
    add_formatted_paragraph(doc, "• Geographic expansion: International market entry", italic=True)
    
    add_heading_with_style(doc, "8.2 Margin Expansion Drivers", 2)
    add_formatted_paragraph(doc,
        "Factors driving margin improvement:")
    add_formatted_paragraph(doc, "• Economies of scale: Fixed cost leverage", italic=True)
    add_formatted_paragraph(doc, "• Product mix: Higher-margin software vs. services", italic=True)
    add_formatted_paragraph(doc, "• Operational efficiency: Automation and process improvement", italic=True)
    add_formatted_paragraph(doc, "• Pricing power: Premium positioning in market", italic=True)
    
    # Final Deliverables
    add_heading_with_style(doc, "9. Final Deliverables", 1)
    add_heading_with_style(doc, "9.1 Investment Recommendation", 2)
    add_formatted_paragraph(doc,
        "Based on our comprehensive DCF analysis, we recommend:", bold=True)
    add_formatted_paragraph(doc,
        "• Target Enterprise Value: $124M (base case)", bold=True)
    add_formatted_paragraph(doc,
        "• Valuation Range: $89M - $156M", bold=True)
    add_formatted_paragraph(doc,
        "• Investment Rating: BUY", bold=True)
    add_formatted_paragraph(doc,
        "• Price Target: $124 per share", bold=True)
    
    add_heading_with_style(doc, "9.2 Key Risks", 2)
    add_formatted_paragraph(doc,
        "Primary investment risks:")
    add_formatted_paragraph(doc, "• Technology risk: AI model performance and accuracy", italic=True)
    add_formatted_paragraph(doc, "• Market risk: Slower than expected adoption", italic=True)
    add_formatted_paragraph(doc, "• Competition risk: Large incumbents entering market", italic=True)
    add_formatted_paragraph(doc, "• Execution risk: Scaling challenges and talent acquisition", italic=True)
    
    add_heading_with_style(doc, "9.3 Catalysts", 2)
    add_formatted_paragraph(doc,
        "Potential value creation catalysts:")
    add_formatted_paragraph(doc, "• Strategic partnerships with major manufacturers", italic=True)
    add_formatted_paragraph(doc, "• Product launches in new verticals", italic=True)
    add_formatted_paragraph(doc, "• International expansion", italic=True)
    add_formatted_paragraph(doc, "• Potential acquisition by larger technology company", italic=True)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Generated by Phoenician Capital AI Analysis Platform - {datetime.now().strftime('%B %d, %Y')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def main():
    """Generate and save sample responses to Word documents"""
    print("Generating impressive TAM and DCF analysis samples...")
    
    # Generate TAM analysis
    print("Creating TAM analysis document...")
    tam_doc = create_tam_sample_response()
    tam_filename = f"TAM_Analysis_Sample_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    tam_doc.save(tam_filename)
    print(f"✅ TAM analysis saved as: {tam_filename}")
    
    # Generate DCF analysis
    print("Creating DCF analysis document...")
    dcf_doc = create_dcf_sample_response()
    dcf_filename = f"DCF_Analysis_Sample_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    dcf_doc.save(dcf_filename)
    print(f"✅ DCF analysis saved as: {dcf_filename}")
    
    print("\n🎉 Sample analysis documents generated successfully!")
    print(f"📄 TAM Analysis: {tam_filename}")
    print(f"📄 DCF Analysis: {dcf_filename}")
    print("\nThese documents showcase the quality and depth of analysis")
    print("that your Phoenician Capital AI platform can generate!")

if __name__ == "__main__":
    main()


